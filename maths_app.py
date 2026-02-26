import streamlit as st
import os
import fitz  # PyMuPDF
import gdown
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import date

ADMIN_PASSWORD = "9709Admin"
# --- 1. GOOGLE DRIVE CONFIGURATION ---
# Your shared Folder ID
GD_FOLDER_ID = "1MnXORHT0jmoqpDie4SjlVoj3l8w6Go0a"

# Local folder names (must match your Drive sub-folders)
FOLDERS = {
    "June QP": "9709_June_qp",
    "Nov QP": "9709_Nov_qp",
    "June MS": "9709_June_ms",
    "Nov MS": "9709_Nov_ms"
}

# Create local folders if they don't exist
for folder in FOLDERS.values():
    if not os.path.exists(folder):
        os.makedirs(folder)

# Function to sync from Google Drive
def sync_from_drive():
    try:
        with st.spinner("🔄 Syncing with Google Drive..."):
            # This downloads the content of your GD folder into the app
            # 'remaining_ok=True' helps avoid errors if folders are empty
            gdown.download_folder(id=GD_FOLDER_ID, output=".", quiet=True, remaining_ok=True)
        st.success("✅ Library Updated from Google Drive!")
    except Exception as e:
        st.error(f"Sync Error: {e}")

# --- 2. APP STATE ---
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []
if 'search_results' not in st.session_state:
    st.session_state.search_results = []

# --- 3. UI LAYOUT ---
st.set_page_config(page_title="9709 Handout Builder", layout="wide")
st.title(" PUSAT TINGKATAN ENAM SENGKURONG ")
st.title("📚 9709 Mathematics PYP Handout Platform")

# Add a Sync Button in the Sidebar
with st.sidebar:
    st.header("Cloud Controls")
    if st.button("🔄 Sync New Files"):
        sync_from_drive()
    st.info("Drop new PDFs into Google Drive, then click Sync here.")

# ... (rest of your existing search and export code below) ...
##################################################################

def get_filename_pattern(month, year, paper_type, paper_code):
    short_year = year[-2:]
    month_code = 's' if month == "June" else 'w'
    return f"9709_{month_code}{short_year}_{paper_type}_{paper_code}"

def search_pdfs(keyword_list, folder_path):
    results = []
    if not os.path.exists(folder_path): return results
    for file in os.listdir(folder_path):
        if file.endswith(".pdf"):
            try:
                doc = fitz.open(os.path.join(folder_path, file))
                for page_num in range(len(doc)):
                    text = doc[page_num].get_text().lower()
                    if all(k.lower() in text for k in keyword_list):
                        results.append({
                            "file": file, "page": page_num, "path": os.path.join(folder_path, file)
                        })
                doc.close()
            except:
                continue
    return results

tab1, tab2, tab3, tab4 = st.tabs(["🔍 Search via Keyword & Extract", "📅 View PYP Question & Answer", "📝 Export Merged Handout", "⚙️ Admin for Upload/Delete PYP"])

# --- TAB 1: SEARCH & COLLECT ---
with tab1:
    st.header("Search for Specific Questions")
    col_input, col_reset = st.columns([4, 1])

    with col_input:
        keywords = st.text_input("Enter keywords (e.g., 'vector')")

    with col_reset:
        st.write(" ")  # Padding
        if st.button("🔄 Clear Search"):
            st.session_state.search_results = []
            st.rerun()

    if st.button("Search Papers", type="primary"):
        if keywords:
            with st.spinner("Scanning all papers..."):
                all_results = []
                for folder_path in FOLDERS.values():
                    all_results += search_pdfs([keywords], folder_path)
                st.session_state.search_results = all_results
        else:
            st.warning("Please enter a keyword first.")

    if st.session_state.search_results:
        st.write(f"Found {len(st.session_state.search_results)} pages:")
        for idx, item in enumerate(st.session_state.search_results):
            c1, c2 = st.columns([4, 1])
            c1.write(f"📄 **{item['file']}** (Page {item['page'] + 1})")
            if c2.button("➕ Add", key=f"add_{idx}"):
                st.session_state.handout_basket.append(item)
                st.toast(f"Added Page {item['page'] + 1}!")

# --- TAB 2: VIEW FULL PAPERS ---
with tab2:
    st.header("Quick Download: Full Papers")
    c1, c2, c3 = st.columns(3)
    with c1:
        v_year = st.selectbox("Year", [str(y) for y in range(2030, 2003, -1)])
    with c2:
        v_month = st.selectbox("Month", ["June", "Nov"])
    with c3:
        v_paper = st.selectbox("Paper Variant", ["13", "33", "43", "53", "63"])

    qp_name = get_filename_pattern(v_month, v_year, "qp", v_paper) + ".pdf"
    ms_name = get_filename_pattern(v_month, v_year, "ms", v_paper) + ".pdf"

    col_q, col_m = st.columns(2)
    with col_q:
        path = os.path.join(FOLDERS[f"{v_month} QP"], qp_name)
        if os.path.exists(path):
            st.success(f"Found QP: {qp_name}")
            with open(path, "rb") as f:
                st.download_button("Download Full QP", f, file_name=qp_name)
        else:
            st.error("Question Paper not found.")

    with col_m:
        path_ms = os.path.join(FOLDERS[f"{v_month} MS"], ms_name)
        if os.path.exists(path_ms):
            st.success(f"Found MS: {ms_name}")
            with open(path_ms, "rb") as f:
                st.download_button("Download Full MS", f, file_name=ms_name)
        else:
            st.error("Mark Scheme not found.")

# --- TAB 3: EXPORT HANDOUT ---
with tab3:
    st.header("Worksheet Export")
    if not st.session_state.handout_basket:
        st.info("Your basket is empty. Add questions from the Search tab.")
    else:
        st.write(f"Items in basket: **{len(st.session_state.handout_basket)}**")
        if st.button("🗑️ Empty Basket"):
            st.session_state.handout_basket = []
            st.rerun()

        if st.button("🪄 Generate Word Handout (.docx)", type="primary"):
            doc = Document()
            doc.add_heading('9709 Mathematics Worksheet', 0)
            for item in st.session_state.handout_basket:
                doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
                pdf_doc = fitz.open(item['path'])
                page = pdf_doc.load_page(item['page'])
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_data = BytesIO(pix.tobytes("png"))
                doc.add_picture(img_data, width=Inches(6))
                doc.add_page_break()
                pdf_doc.close()

            target = "9709_Custom_Handout.docx"
            doc.save(target)
            with open(target, "rb") as f:
                st.download_button("📥 Click to Download Document", f, file_name=target)

# --- TAB 4: ADMIN ---
with tab4:
    st.header("Admin Management")
    pwd = st.text_input("Admin Password", type="password")
    if pwd == ADMIN_PASSWORD:
        u_col, d_col = st.columns(2)
        with u_col:
            st.subheader("Upload")
            dest = st.selectbox("Destination", list(FOLDERS.keys()))
            up_files = st.file_uploader("Select PDFs", type="pdf", accept_multiple_files=True)
            if st.button("Upload"):
                for f in up_files:
                    with open(os.path.join(FOLDERS[dest], f.name), "wb") as s: s.write(f.getbuffer())
                st.success("Uploaded!")
        with d_col:
            st.subheader("Delete")
            dest_d = st.selectbox("Clean Folder", list(FOLDERS.keys()))
            to_del = st.selectbox("File to Remove", ["---"] + os.listdir(FOLDERS[dest_d]))
            if to_del != "---" and st.button("Delete Permanently"):
                os.remove(os.path.join(FOLDERS[dest_d], to_del))
                st.rerun()

# --- FOOTER ---
st.markdown("---")

# Using a single container with centered alignment
st.markdown(
    """
    <div style="text-align: center; width: 100%;">
        <p style="font-size: 20px; font-weight: bold; margin-bottom: 5px;">
            ✨ PTES 9709 Mathematics Resource Portal ✨
        </p>
        <p style="font-size: 16px; font-weight: bold; letter-spacing: 0.5px;">
            <span style="color: #FF0000;">🔴 Academically Excellence</span> | 
            <span style="color: #FFD700;">🟡 Future Readiness</span> | 
            <span style="color: #0070FF;">🔵 Digital & Integrity</span> | 
            <span style="color: #28A745;">🟢 Holistic & Growth</span>
        </p>
        <p style="color: gray; font-size: 14px; margin-top: 10px;">
            Creator: Miss Hajah Nurul Haziqah HN (PTES CS Tutor)
        </p>
    </div>
    """,
    unsafe_allow_html=True
)